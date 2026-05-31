"""
ReportForge DOCX export — generates a .docx from a Report using python-docx.

Soft dependency: python-docx must be installed separately.
    pip install python-docx

Falls back gracefully with ImportError if not available.
"""

from __future__ import annotations

import io
import logging
from typing import Any, TYPE_CHECKING

if TYPE_CHECKING:
	from sqlalchemy.orm import Session
	from .models import Report

log = logging.getLogger(__name__)

try:
	from docx import Document
	from docx.shared import Inches, Pt, RGBColor, Mm
	from docx.enum.text import WD_ALIGN_PARAGRAPH
	from docx.enum.table import WD_TABLE_ALIGNMENT
	from docx.oxml.ns import qn
	from docx.oxml import OxmlElement
	HAS_DOCX = True
except ImportError:
	HAS_DOCX = False
	log.debug("python-docx not installed — DOCX export disabled")


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
	"""Convert #rrggbb to (r, g, b)."""
	h = hex_color.lstrip("#")
	if len(h) == 3:
		h = "".join(c * 2 for c in h)
	return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _apply_branding_header(doc: Any, report: Report) -> None:
	"""Insert company name, logo, and primary color rule in the DOCX header."""
	section = doc.sections[0]
	header = section.header
	header.is_linked_to_previous = False

	p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
	p.clear()

	if report.company_name:
		run = p.add_run(report.company_name)
		run.bold = True
		r, g, b = _hex_to_rgb(report.primary_color or "#003366")
		run.font.color.rgb = RGBColor(r, g, b)
		run.font.size = Pt(14)

	p.alignment = WD_ALIGN_PARAGRAPH.LEFT

	# Horizontal rule below header
	pPr = p._p.get_or_add_pPr()
	pBdr = OxmlElement("w:pBdr")
	bottom = OxmlElement("w:bottom")
	bottom.set(qn("w:val"), "single")
	bottom.set(qn("w:sz"), "6")
	bottom.set(qn("w:space"), "1")
	bottom.set(qn("w:color"), (report.primary_color or "#003366").lstrip("#"))
	pBdr.append(bottom)
	pPr.append(pBdr)


def _apply_branding_footer(doc: Any, report: Report) -> None:
	"""Insert page number and custom footer text."""
	section = doc.sections[0]
	footer = section.footer
	footer.is_linked_to_previous = False

	p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
	p.clear()
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER

	if report.custom_footer_html:
		# Strip HTML tags for DOCX (plain text fallback)
		import re
		plain = re.sub(r"<[^>]+>", "", report.custom_footer_html).strip()
		if plain:
			run = p.add_run(plain + "   ")
			run.font.size = Pt(8)
			run.font.color.rgb = RGBColor(100, 100, 100)

	# Page number field
	run = p.add_run("Page ")
	run.font.size = Pt(8)
	fldChar1 = OxmlElement("w:fldChar")
	fldChar1.set(qn("w:fldCharType"), "begin")
	instrText = OxmlElement("w:instrText")
	instrText.text = "PAGE"
	fldChar2 = OxmlElement("w:fldChar")
	fldChar2.set(qn("w:fldCharType"), "end")
	run._r.append(fldChar1)
	run._r.append(instrText)
	run._r.append(fldChar2)


def generate_docx(report: Report, rows: list[dict[str, Any]]) -> bytes:
	"""
	Render a Report to a DOCX byte string.

	Args:
	    report: the Report ORM instance (with branding fields populated).
	    rows: list of dicts from the report data source query.

	Returns:
	    Raw DOCX bytes suitable for sending as a file download.

	Raises:
	    ImportError: if python-docx is not installed.
	"""
	if not HAS_DOCX:
		raise ImportError(
			"python-docx is required for DOCX export. "
			"Install it with: pip install python-docx"
		)

	doc = Document()

	# ── Page setup ────────────────────────────────────────────────────────
	from docx.shared import Mm as DocxMm
	section = doc.sections[0]
	page_cfg = report.page_config or {}
	margin_top    = page_cfg.get("margin_top_mm",    20)
	margin_bottom = page_cfg.get("margin_bottom_mm", 20)
	margin_left   = page_cfg.get("margin_left_mm",   25)
	margin_right  = page_cfg.get("margin_right_mm",  25)
	section.top_margin    = DocxMm(margin_top)
	section.bottom_margin = DocxMm(margin_bottom)
	section.left_margin   = DocxMm(margin_left)
	section.right_margin  = DocxMm(margin_right)

	# Portrait / landscape
	if report.orientation and report.orientation.value == "landscape":
		from docx.shared import Mm as M
		section.orientation = 1  # WD_ORIENT.LANDSCAPE
		section.page_width, section.page_height = M(297), M(210)
	else:
		from docx.shared import Mm as M
		section.page_width, section.page_height = M(210), M(297)

	# ── Branding: header + footer ─────────────────────────────────────────
	_apply_branding_header(doc, report)
	_apply_branding_footer(doc, report)

	# ── Report title ──────────────────────────────────────────────────────
	title_p = doc.add_heading(report.name, level=1)
	r, g, b = _hex_to_rgb(report.primary_color or "#003366")
	for run in title_p.runs:
		run.font.color.rgb = RGBColor(r, g, b)

	if report.description:
		desc_p = doc.add_paragraph(report.description)
		desc_p.runs[0].italic = True if desc_p.runs else None

	# ── Data table ────────────────────────────────────────────────────────
	if rows:
		columns = list(rows[0].keys())
		table = doc.add_table(rows=1, cols=len(columns))
		table.style = "Table Grid"
		table.alignment = WD_TABLE_ALIGNMENT.CENTER

		# Header row
		hdr_cells = table.rows[0].cells
		primary_r, primary_g, primary_b = _hex_to_rgb(report.primary_color or "#003366")
		for i, col_name in enumerate(columns):
			cell = hdr_cells[i]
			cell.text = col_name.replace("_", " ").title()
			# Bold + color
			for para in cell.paragraphs:
				for run in para.runs:
					run.bold = True
					run.font.color.rgb = RGBColor(255, 255, 255)
				para.alignment = WD_ALIGN_PARAGRAPH.CENTER
			# Background color (primary)
			tc = cell._tc
			tcPr = tc.get_or_add_tcPr()
			shd = OxmlElement("w:shd")
			shd.set(qn("w:val"), "clear")
			shd.set(qn("w:color"), "auto")
			hex_color = f"{primary_r:02X}{primary_g:02X}{primary_b:02X}"
			shd.set(qn("w:fill"), hex_color)
			tcPr.append(shd)

		# Data rows (alternate shading)
		for row_idx, row_data in enumerate(rows):
			row_cells = table.add_row().cells
			for i, col_name in enumerate(columns):
				val = row_data.get(col_name, "")
				row_cells[i].text = "" if val is None else str(val)
				row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
				# Alternate row background
				if row_idx % 2 == 1:
					tc = row_cells[i]._tc
					tcPr = tc.get_or_add_tcPr()
					shd = OxmlElement("w:shd")
					shd.set(qn("w:val"), "clear")
					shd.set(qn("w:color"), "auto")
					shd.set(qn("w:fill"), "F0F4F8")
					tcPr.append(shd)
	else:
		doc.add_paragraph("No data returned by this report.")

	# ── Watermark ─────────────────────────────────────────────────────────
	if report.watermark_text:
		# DOCX watermarks go in the header XML as a shape
		_add_docx_watermark(doc, report.watermark_text, report.watermark_opacity or 0.08)

	buf = io.BytesIO()
	doc.save(buf)
	buf.seek(0)
	return buf.read()


def _add_docx_watermark(doc: Any, text: str, opacity: float) -> None:
	"""Add a diagonal watermark to all pages via the DOCX header."""
	# Watermarks in DOCX are VML shapes in the header — complex XML
	from lxml import etree
	section = doc.sections[0]
	header = section.header
	if not header.paragraphs:
		header.add_paragraph()
	p = header.paragraphs[0]._p

	# Build VML watermark XML
	opacity_int = max(0, min(65536, int((1.0 - opacity) * 65536)))
	ns = {
		"v": "urn:schemas-microsoft-com:vml",
		"o": "urn:schemas-microsoft-com:office:office",
		"w10": "urn:schemas-microsoft-com:office:word",
	}
	pict = etree.SubElement(p, "{urn:schemas-microsoft-com:vml}shape", attrib={
		"id":   "watermark",
		"type": "#_x0000_t136",
		"style": (
			"position:absolute;left:0;top:0;width:100%;height:100%;"
			"z-index:-251658240;mso-position-horizontal:center;"
			"mso-position-horizontal-relative:margin;"
			"mso-position-vertical:center;"
			"mso-position-vertical-relative:margin;"
			"rotation:315;"
		),
		"fillcolor": "#d0d0d0",
		"stroked": "f",
	})
	txbx = etree.SubElement(pict, "{urn:schemas-microsoft-com:vml}textbox")
	txbx.set("style", "mso-fit-shape-to-text:t")
	txBody = etree.SubElement(txbx, "txbxContent")
	wml_p = etree.SubElement(txBody, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p")
	wml_r = etree.SubElement(wml_p, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
	wml_rPr = etree.SubElement(wml_r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr")
	wml_rFonts = etree.SubElement(
		wml_rPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
	)
	wml_rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii", "Calibri")
	wml_sz = etree.SubElement(wml_rPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sz")
	wml_sz.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "288")
	wml_color = etree.SubElement(wml_rPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color")
	wml_color.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "D0D0D0")
	wml_t = etree.SubElement(wml_r, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
	wml_t.text = text.upper()
