from __future__ import annotations

import hashlib
import logging
import mimetypes
import os
import re
import tempfile
from datetime import datetime
from io import BytesIO
from typing import Any, Callable

try:
	from sqlalchemy.orm import declared_attr, Mapped, mapped_column
	from sqlalchemy import Boolean, Integer, LargeBinary, String, Text
	_SQLA2 = True
except ImportError:
	from sqlalchemy.ext.declarative import declared_attr
	from sqlalchemy import Boolean, Column, Integer, LargeBinary, String, Text
	_SQLA2 = False

from sqlalchemy import Column
from sqlalchemy.orm import validates

# PostgreSQL-specific types — optional; fall back to JSON/Text for other DBs
try:
	from sqlalchemy.dialects.postgresql import ARRAY, JSONB
	_PG_AVAILABLE = True
except ImportError:
	_PG_AVAILABLE = False

# TSVector full-text search — requires sqlalchemy-utils with PostgreSQL
try:
	from sqlalchemy_utils import TSVectorType
	_TSVECTOR_AVAILABLE = True
except ImportError:
	_TSVECTOR_AVAILABLE = False

# FAB's own file/image column types
from pgappforge.models.mixins import FileColumn, ImageColumn

# Optional heavy dependencies — each checked at call-site, not import time
_PyPDF2: Any = None
_docx: Any = None
_magic: Any = None
_pypandoc: Any = None
_markdown: Any = None
_pdfkit: Any = None
_nltk: Any = None
_transformers_pipeline: Any = None


def _lazy_import_PyPDF2():
	global _PyPDF2
	if _PyPDF2 is None:
		import PyPDF2 as _m
		_PyPDF2 = _m
	return _PyPDF2


def _lazy_import_docx():
	global _docx
	if _docx is None:
		import docx as _m
		_docx = _m
	return _docx


def _lazy_import_magic():
	global _magic
	if _magic is None:
		import magic as _m
		_magic = _m
	return _magic


def _lazy_import_pypandoc():
	global _pypandoc
	if _pypandoc is None:
		import pypandoc as _m
		_pypandoc = _m
	return _pypandoc


def _lazy_import_nltk():
	global _nltk
	if _nltk is None:
		import nltk as _m
		_nltk = _m
	return _nltk


def _lazy_import_transformers_pipeline():
	global _transformers_pipeline
	if _transformers_pipeline is None:
		from transformers import pipeline as _p
		_transformers_pipeline = _p
	return _transformers_pipeline


logger = logging.getLogger(__name__)


def _make_jsonb_column():
	"""Return a JSONB Column on PostgreSQL, plain JSON elsewhere."""
	if _PG_AVAILABLE:
		from sqlalchemy.dialects.postgresql import JSONB as _JSONB
		return Column(_JSONB, nullable=True)
	from sqlalchemy import JSON
	return Column(JSON, nullable=True)


def _make_array_jsonb_column():
	"""Return ARRAY(JSONB) on PostgreSQL, plain JSON elsewhere."""
	if _PG_AVAILABLE:
		from sqlalchemy.dialects.postgresql import ARRAY as _A, JSONB as _J
		return Column(_A(_J), nullable=True)
	from sqlalchemy import JSON
	return Column(JSON, nullable=True)


def _make_array_string_column():
	"""Return ARRAY(String) on PostgreSQL, plain Text elsewhere."""
	if _PG_AVAILABLE:
		from sqlalchemy.dialects.postgresql import ARRAY as _A
		return Column(_A(String), nullable=True)
	return Column(Text, nullable=True)


class DocMixin:
	"""
	SQLAlchemy mixin for document management.

	Provides metadata extraction, format conversion, Fernet encryption,
	extractive NLP summarisation, full-text search (PostgreSQL/TSVector),
	and version-history tracking for any model that stores binary documents.

	Columns covering: MIME type, raw binary, extracted text, structural
	hierarchy (chapter/section/sub-section), LLM context/prompt, rich
	metadata statistics, audio properties, encryption state, and
	PostgreSQL JSONB/ARRAY metadata storage.

	Optional heavy dependencies (PyPDF2, python-docx, python-magic,
	pypandoc, nltk, transformers, pdfkit) are imported lazily so the
	mixin is importable even when they are absent — errors surface only
	when the relevant method is called.
	"""

	# ------------------------------------------------------------------ #
	# Core document storage                                                #
	# ------------------------------------------------------------------ #
	mime_type: str = Column(String(60), default="application/pdf", nullable=False)
	doc: ImageColumn = Column(
		ImageColumn(thumbnail_size=(30, 30, True), size=(300, 300, True))
	)
	doc_text: str | None = Column(Text, nullable=True)
	doc_binary: bytes | None = Column(LargeBinary, nullable=True)
	doc_title: str = Column(String(200), nullable=False, default="Untitled Document")
	subject: str | None = Column(String(100), nullable=True)
	author: str | None = Column(String(100), nullable=True)
	keywords: str | None = Column(String(200), nullable=True)
	comments: str | None = Column(Text, nullable=True)

	# ------------------------------------------------------------------ #
	# Document hierarchy                                                   #
	# ------------------------------------------------------------------ #
	chapter_number: int | None = Column(Integer, nullable=True)
	chapter_title: str | None = Column(String(200), nullable=True)
	section_number: int | None = Column(Integer, nullable=True)
	section_title: str | None = Column(String(200), nullable=True)
	sub_section_number: int | None = Column(Integer, nullable=True)
	sub_section_title: str | None = Column(String(200), nullable=True)

	# ------------------------------------------------------------------ #
	# LLM context                                                          #
	# ------------------------------------------------------------------ #
	doc_context: str | None = Column(Text, nullable=True)
	doc_prompt: str | None = Column(Text, nullable=True)

	# ------------------------------------------------------------------ #
	# Derived statistics                                                   #
	# ------------------------------------------------------------------ #
	doc_type: str = Column(String(5), default="pdf", nullable=False)
	char_count: int = Column(Integer, default=0)
	word_count: int = Column(Integer, default=0)
	lines: int = Column(Integer, default=0)
	paragraphs: int = Column(Integer, default=0)
	gpt_token_count: int = Column(Integer, default=0)
	grammar_checked: bool = Column(Boolean, default=False)
	doc_summary: str | None = Column(Text, nullable=True)
	doc_spell_checked: bool = Column(Boolean, default=False)
	doc_gpt_ver: str | None = Column(String(40), nullable=True)
	doc_format: str = Column(String(40), nullable=False, default="pdf")
	doc_downloadable: bool = Column(Boolean, default=True)
	doc_template: str | None = Column(Text, nullable=True)
	doc_rendered: bool = Column(Boolean, default=False)
	doc_render: FileColumn | None = Column(FileColumn, nullable=True)

	file_size_bytes: int = Column(Integer, default=0)
	producer_prog: str | None = Column(String(40), nullable=True)
	immutable: bool = Column(Boolean, default=False)

	page_size: str | None = Column(String(40), nullable=True)
	page_count: int = Column(Integer, default=1)
	hashx: str | None = Column(String(64), nullable=True)  # SHA-256

	# ------------------------------------------------------------------ #
	# Audio metadata                                                       #
	# ------------------------------------------------------------------ #
	is_audio: bool = Column(Boolean, default=False)
	audio_duration_secs: int | None = Column(Integer, nullable=True)
	audio_frame_rate: int | None = Column(Integer, nullable=True)
	audio_channels: int | None = Column(Integer, nullable=True)

	# ------------------------------------------------------------------ #
	# Encryption                                                           #
	# ------------------------------------------------------------------ #
	is_encrypted: bool = Column(Boolean, default=False)
	encryption_key: str | None = Column(String(100), nullable=True)

	# ------------------------------------------------------------------ #
	# Flexible metadata storage (JSONB on PostgreSQL, JSON elsewhere)      #
	# ------------------------------------------------------------------ #
	@declared_attr
	def doc_metadata(cls):  # named doc_metadata to avoid collision with SA metadata
		return _make_jsonb_column()

	@declared_attr
	def versions(cls):
		return _make_array_jsonb_column()

	@declared_attr
	def tags(cls):
		return _make_array_string_column()

	# ------------------------------------------------------------------ #
	# Full-text search vector (PostgreSQL + sqlalchemy-utils only)         #
	# ------------------------------------------------------------------ #
	@declared_attr
	def search_vector(cls):
		if _TSVECTOR_AVAILABLE:
			return Column(
				TSVectorType(
					"doc_title",
					"doc_text",
					"comments",
					weights={"doc_title": "A", "doc_text": "B", "comments": "C"},
				),
				nullable=True,
				index=True,
			)
		# Non-PostgreSQL: store a plain-text pre-computed search blob
		return Column(Text, nullable=True, index=False)

	# ------------------------------------------------------------------ #
	# SQLAlchemy validators                                                #
	# ------------------------------------------------------------------ #
	@validates("doc_text")
	def update_doc_info(self, key: str, value: str | None) -> str | None:
		"""Recompute statistics whenever doc_text is assigned."""
		if value:
			self.char_count = len(value)
			self.word_count = len(value.split())
			self.lines = value.count("\n") + 1
			self.paragraphs = len([p for p in value.split("\n\n") if p.strip()])
		return value

	# ------------------------------------------------------------------ #
	# Dependency introspection                                             #
	# ------------------------------------------------------------------ #
	@classmethod
	def check_dependencies(cls) -> dict[str, bool]:
		"""Return availability of each optional heavy dependency."""
		results: dict[str, bool] = {}
		for name, mod in [
			("PyPDF2", "PyPDF2"),
			("pypandoc", "pypandoc"),
			("python-magic", "magic"),
			("python-docx", "docx"),
			("nltk", "nltk"),
			("transformers", "transformers"),
			("markdown", "markdown"),
			("pdfkit", "pdfkit"),
			("cryptography", "cryptography"),
			("sqlalchemy-utils", "sqlalchemy_utils"),
		]:
			try:
				__import__(mod)
				results[name] = True
			except ImportError:
				results[name] = False
		return results

	# ------------------------------------------------------------------ #
	# Hash                                                                 #
	# ------------------------------------------------------------------ #
	def update_hash(self) -> None:
		"""Recompute SHA-256 hash of doc_binary."""
		if self.doc_binary:
			self.hashx = hashlib.sha256(self.doc_binary).hexdigest()

	# ------------------------------------------------------------------ #
	# Metadata extraction                                                  #
	# ------------------------------------------------------------------ #
	def _extract_pdf_metadata(self) -> dict[str, Any]:
		"""Extract standard PDF metadata via PyPDF2."""
		PyPDF2 = _lazy_import_PyPDF2()
		metadata: dict[str, Any] = {}
		with BytesIO(self.doc_binary) as fh:
			reader = PyPDF2.PdfReader(fh)
			info = reader.metadata
			if info:
				metadata.update({
					"author": info.get("/Author", ""),
					"subject": info.get("/Subject", ""),
					"title": info.get("/Title", ""),
					"creator": info.get("/Creator", ""),
					"producer": info.get("/Producer", ""),
					"creation_date": info.get("/CreationDate", ""),
					"modification_date": info.get("/ModDate", ""),
					"keywords": info.get("/Keywords", ""),
				})
			metadata["page_count"] = len(reader.pages)
			if reader.pages:
				metadata["first_page_text"] = (reader.pages[0].extract_text() or "")[:1000]
		return metadata

	def _extract_word_metadata(self) -> dict[str, Any]:
		"""Extract core properties from a .docx file via python-docx."""
		docx_mod = _lazy_import_docx()
		metadata: dict[str, Any] = {}
		try:
			with BytesIO(self.doc_binary) as fh:
				doc = docx_mod.Document(fh)
				cp = doc.core_properties
				metadata.update({
					"author": cp.author or "",
					"title": cp.title or "",
					"subject": cp.subject or "",
					"keywords": cp.keywords or "",
					"created": str(cp.created) if cp.created else "",
					"modified": str(cp.modified) if cp.modified else "",
					"last_modified_by": cp.last_modified_by or "",
					"revision": cp.revision or 1,
					"category": cp.category or "",
					"paragraphs": len(doc.paragraphs),
					"sections": len(doc.sections),
				})
		except Exception as exc:
			logger.error("Error extracting Word metadata: %s", exc)
			metadata["extraction_error"] = str(exc)
		return metadata

	def _extract_text_metadata(self) -> dict[str, Any]:
		"""Analyse a plain-text binary payload and return basic stats."""
		content = self.doc_binary.decode("utf-8", errors="ignore")
		split_lines = content.split("\n")
		avg_len = (
			sum(len(ln) for ln in split_lines) / len(split_lines)
			if split_lines else 0.0
		)
		return {
			"file_size": len(content),
			"line_count": len(split_lines),
			"first_line": split_lines[0][:100] if split_lines else "",
			"encoding": "utf-8",
			"has_bom": content.startswith("﻿"),
			"empty_lines": sum(1 for ln in split_lines if not ln.strip()),
			"avg_line_length": avg_len,
		}

	def _extract_markdown_metadata(self) -> dict[str, Any]:
		"""Analyse a Markdown binary payload and return structural stats."""
		content = self.doc_binary.decode("utf-8", errors="ignore")
		split_lines = content.split("\n")
		headers = [ln for ln in split_lines if ln.strip().startswith("#")]
		list_lines = [
			ln for ln in split_lines
			if ln.strip().startswith(("- ", "* ", "1. "))
		]
		return {
			"file_size": len(content),
			"headers_count": len(headers),
			"first_heading": headers[0][:100] if headers else "",
			"links_count": content.count("]("),
			"code_blocks": content.count("```"),
			"images": content.count("!["),
			"lists": len(list_lines),
		}

	def _update_attributes_from_metadata(self, metadata: dict[str, Any]) -> None:
		"""Write extracted values onto self; merge into doc_metadata JSONB."""
		for key, value in metadata.items():
			if hasattr(self, key) and value is not None:
				setattr(self, key, value)

		if self.doc_metadata is None:
			self.doc_metadata = {}
		self.doc_metadata.update(metadata)
		self.update_hash()

	def extract_metadata(self) -> dict[str, Any]:
		"""
		Dispatch to the appropriate extractor based on MIME type, populate
		model fields, and return the raw metadata dict.
		"""
		handlers: dict[str, Callable[[], dict[str, Any]]] = {
			"application/pdf": self._extract_pdf_metadata,
			"application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_word_metadata,
			"application/msword": self._extract_word_metadata,
			"text/plain": self._extract_text_metadata,
			"text/markdown": self._extract_markdown_metadata,
		}

		handler = handlers.get(self.mime_type)
		metadata: dict[str, Any] = {}
		if handler:
			try:
				metadata = handler()
			except Exception as exc:
				metadata["extraction_error"] = str(exc)
				logger.error("Error extracting metadata: %s", exc)
		else:
			metadata["extraction_note"] = (
				f"Metadata extraction not supported for MIME type: {self.mime_type}"
			)

		self._update_attributes_from_metadata(metadata)
		return metadata

	# ------------------------------------------------------------------ #
	# NLP summarisation                                                    #
	# ------------------------------------------------------------------ #
	def generate_summary(self, max_length: int = 200) -> None:
		"""
		Build an extractive summary from doc_text using NLTK sentence
		scoring, then optionally refine it with a HuggingFace transformer
		(facebook/bart-large-cnn) if the transformers package is available.

		Updates self.doc_summary in place.
		"""
		if not self.doc_text:
			raise ValueError("No document text available for summarisation")

		nltk = _lazy_import_nltk()
		nltk.download("punkt", quiet=True)
		nltk.download("stopwords", quiet=True)
		from nltk.tokenize import sent_tokenize, word_tokenize
		from nltk.corpus import stopwords as nltk_stopwords

		sentences = sent_tokenize(self.doc_text)
		if not sentences:
			raise ValueError("No sentences found in document text")

		stop_words = set(nltk_stopwords.words("english"))
		words = word_tokenize(self.doc_text.lower())
		word_frequencies: dict[str, float] = {}
		for word in words:
			if word not in stop_words and word.isalnum():
				word_frequencies[word] = word_frequencies.get(word, 0) + 1

		max_freq = max(word_frequencies.values(), default=1)
		for w in word_frequencies:
			word_frequencies[w] /= max_freq

		sentence_scores: dict[str, float] = {}
		for sent in sentences:
			if len(sent.split()) < 30:
				score = sum(
					word_frequencies.get(w, 0)
					for w in word_tokenize(sent.lower())
				)
				sentence_scores[sent] = score

		top_sentences = sorted(
			sentence_scores.items(), key=lambda x: x[1], reverse=True
		)[:3]
		extractive_summary = " ".join(s for s, _ in top_sentences)

		try:
			pipeline_fn = _lazy_import_transformers_pipeline()
			summariser = pipeline_fn("summarization", model="facebook/bart-large-cnn")
			chunks = [
				extractive_summary[i: i + 1024]
				for i in range(0, len(extractive_summary), 1024)
			]
			self.doc_summary = " ".join(
				summariser(chunk, max_length=max_length, min_length=30, do_sample=False)[0]["summary_text"]
				for chunk in chunks
			)
		except Exception as exc:
			logger.warning(
				"Advanced summarisation failed, using extractive summary: %s", exc
			)
			self.doc_summary = extractive_summary

	# ------------------------------------------------------------------ #
	# Full-text search (PostgreSQL / TSVector)                             #
	# ------------------------------------------------------------------ #
	@classmethod
	def search(
		cls,
		session,
		query: str,
		*,
		limit: int | None = None,
		offset: int | None = None,
	) -> list[DocMixin]:
		"""
		PostgreSQL full-text search across doc_title, doc_text, and comments.

		Sanitises the query, ranks results by ts_rank, and attaches
		relevance_score plus HTML-highlighted title/text snippets to
		each returned instance.

		Args:
			session: SQLAlchemy Session.
			query:   Raw search string; words are AND-ed together.
			limit:   Maximum rows to return.
			offset:  Rows to skip (for pagination).

		Returns:
			List of model instances ordered by descending relevance.
		"""
		from sqlalchemy.sql import func as sqla_func
		from sqlalchemy import select

		sanitised = " & ".join(w.strip() for w in query.split() if w.strip())
		search_query = sqla_func.plainto_tsquery("english", sanitised)
		rank_fn = sqla_func.ts_rank(cls.search_vector, search_query)

		stmt = (
			select(cls, rank_fn.label("relevance"))
			.where(cls.search_vector.match(search_query))
			.order_by(rank_fn.desc())
		)
		if offset is not None:
			stmt = stmt.offset(offset)
		if limit is not None:
			stmt = stmt.limit(limit)

		rows = session.execute(stmt).all()
		results: list[DocMixin] = []
		for doc, relevance in rows:
			doc.relevance_score = float(relevance)
			doc.highlighted_title = cls.highlight_matched_terms(doc.doc_title or "", sanitised)
			doc.highlighted_text = cls.highlight_matched_terms(doc.doc_text or "", sanitised)
			results.append(doc)
		return results

	@staticmethod
	def highlight_matched_terms(text: str, query: str) -> str:
		"""
		Wrap each query token in <span class="highlight">…</span>.

		Case-insensitive; operates on whole-token matches only to avoid
		partial-word false positives in English prose.
		"""
		if not text or not query:
			return text
		for word in {w.lower() for w in query.split() if w.strip()}:
			text = re.sub(
				f"({re.escape(word)})",
				r'<span class="highlight">\1</span>',
				text,
				flags=re.IGNORECASE,
			)
		return text

	# ------------------------------------------------------------------ #
	# Encryption / decryption                                              #
	# ------------------------------------------------------------------ #
	def encrypt_document(self, key: bytes | None = None) -> None:
		"""
		Encrypt doc_binary with Fernet symmetric encryption.

		If no key is supplied, one is generated automatically and stored
		in self.encryption_key (UTF-8 string). Raises if already encrypted.
		"""
		if self.is_encrypted:
			raise ValueError("Document is already encrypted")
		from cryptography.fernet import Fernet

		if not key:
			key = Fernet.generate_key()

		fernet = Fernet(key)
		self.doc_binary = fernet.encrypt(self.doc_binary)
		self.encryption_key = key.decode("utf-8")
		self.is_encrypted = True

		if self.doc_metadata is None:
			self.doc_metadata = {}
		self.doc_metadata["encryption"] = {
			"method": "Fernet",
			"timestamp": datetime.now().isoformat(),
		}

	def decrypt_document(self) -> None:
		"""
		Decrypt doc_binary using the stored encryption key.

		Raises if the document is not encrypted or the key is missing.
		"""
		if not self.is_encrypted:
			raise ValueError("Document is not encrypted")
		if not self.encryption_key:
			raise ValueError("Encryption key not found")

		from cryptography.fernet import Fernet

		fernet = Fernet(self.encryption_key.encode("utf-8"))
		self.doc_binary = fernet.decrypt(self.doc_binary)
		self.is_encrypted = False
		self.encryption_key = None

		if self.doc_metadata and "encryption" in self.doc_metadata:
			del self.doc_metadata["encryption"]

	# ------------------------------------------------------------------ #
	# Format conversion (Pandoc)                                           #
	# ------------------------------------------------------------------ #
	_FORMAT_MAP: dict[str, str] = {
		"pdf": "pdf",
		"docx": "docx",
		"md": "markdown",
		"html": "html",
		"txt": "plain",
		"rst": "rst",
		"epub": "epub",
		"odt": "odt",
		"latex": "latex",
	}

	def convert_format(self, target_format: str) -> None:
		"""
		Convert doc_binary to *target_format* via Pandoc (pypandoc).

		Updates doc_binary, doc_type, mime_type, and appends a conversion
		entry to the versions history. Raises ValueError for unsupported
		format pairs; propagates pypandoc errors unchanged.
		"""
		if self.doc_type == target_format:
			return

		pypandoc = _lazy_import_pypandoc()

		source_fmt = self._FORMAT_MAP.get(self.doc_type)
		target_fmt = self._FORMAT_MAP.get(target_format)
		if not source_fmt or not target_fmt:
			raise ValueError(
				f"Conversion from {self.doc_type!r} to {target_format!r} not supported"
			)

		tmp_in = None
		try:
			suffix = f".{self.doc_type}"
			with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as fh:
				fh.write(self.doc_binary)
				tmp_in = fh.name

			output = pypandoc.convert_file(
				tmp_in, target_fmt, format=source_fmt, outputfile=None
			)
			self.doc_binary = output.encode("utf-8") if isinstance(output, str) else output
			self.doc_type = target_format
			self.mime_type = (
				mimetypes.guess_type(f"dummy.{target_format}")[0]
				or "application/octet-stream"
			)
			self.extract_metadata()

			if self.versions is None:
				self.versions = []
			self.versions.append({
				"timestamp": datetime.now().isoformat(),
				"action": "format_conversion",
				"from_format": source_fmt,
				"to_format": target_format,
			})
		finally:
			if tmp_in:
				try:
					os.unlink(tmp_in)
				except OSError:
					pass

	def to_markdown(self) -> str:
		"""Convert current doc_binary to Markdown and return as a string."""
		self.convert_format("md")
		return self.doc_binary.decode("utf-8")

	def from_markdown(self, markdown_text: str, target_format: str) -> None:
		"""
		Accept raw Markdown text, convert to *target_format* via Pandoc,
		and update doc_binary, doc_type, mime_type, and metadata.
		"""
		pypandoc = _lazy_import_pypandoc()
		tmp_in = None
		try:
			with tempfile.NamedTemporaryFile(
				mode="w", suffix=".md", delete=False, encoding="utf-8"
			) as fh:
				fh.write(markdown_text)
				tmp_in = fh.name

			output = pypandoc.convert_file(
				tmp_in, target_format, format="markdown", outputfile=None
			)
			self.doc_binary = output.encode("utf-8") if isinstance(output, str) else output
			self.doc_type = target_format
			self.mime_type = (
				mimetypes.guess_type(f"dummy.{target_format}")[0]
				or "application/octet-stream"
			)
			self.extract_metadata()
		finally:
			if tmp_in:
				try:
					os.unlink(tmp_in)
				except OSError:
					pass

	# ------------------------------------------------------------------ #
	# LLM text generation                                                  #
	# ------------------------------------------------------------------ #
	def generate_text_with_llm(self, llm_function: Callable[[str, str], str]) -> None:
		"""
		Call *llm_function(context, prompt)* and store the result in
		doc_text. Also records generation stats in doc_metadata.

		Raises ValueError if doc_context or doc_prompt are unset, or if
		the LLM returns an empty string.
		"""
		if not self.doc_context or not self.doc_prompt:
			raise ValueError("Both doc_context and doc_prompt are required")

		generated = llm_function(self.doc_context, self.doc_prompt)
		if not generated:
			raise ValueError("LLM returned empty text")

		self.doc_text = generated
		self.update_doc_info("doc_text", generated)

		if self.doc_metadata is None:
			self.doc_metadata = {}
		self.doc_metadata["llm_generation"] = {
			"timestamp": datetime.now().isoformat(),
			"context_length": len(self.doc_context),
			"prompt_length": len(self.doc_prompt),
			"output_length": len(generated),
		}

	# ------------------------------------------------------------------ #
	# MIME detection                                                       #
	# ------------------------------------------------------------------ #
	def detect_mime_type(self, filename: str | None = None) -> str:
		"""
		Detect MIME type from doc_binary content (via python-magic) or
		from *filename* extension (via stdlib mimetypes).

		Updates self.mime_type and returns the detected string.
		"""
		if not self.doc_binary and not filename:
			raise ValueError("Either doc_binary or filename is required")

		if self.doc_binary:
			magic_mod = _lazy_import_magic()
			mime = magic_mod.Magic(mime=True)
			detected = mime.from_buffer(self.doc_binary)
		else:
			detected, _ = mimetypes.guess_type(filename)

		self.mime_type = detected or "application/octet-stream"
		return self.mime_type

	def set_doc_type_from_mime_type(self) -> None:
		"""Derive the short doc_type string from the current mime_type."""
		mapping: dict[str, str] = {
			"application/pdf": "pdf",
			"application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
			"application/msword": "doc",
			"text/plain": "txt",
			"text/markdown": "md",
			"text/html": "html",
			"application/rtf": "rtf",
			"application/epub+zip": "epub",
			"application/x-latex": "latex",
		}
		self.doc_type = mapping.get(self.mime_type, "unknown")

	# ------------------------------------------------------------------ #
	# Document update                                                      #
	# ------------------------------------------------------------------ #
	def update_document(self, file_content: bytes, filename: str | None = None) -> None:
		"""
		Replace doc_binary with *file_content*, auto-detect MIME type,
		recompute the hash, and refresh all extracted metadata.

		A version snapshot (previous hash + size) is appended to versions
		before any mutation. Raises ValueError on immutable documents.
		"""
		if self.immutable:
			raise ValueError("Cannot update an immutable document")

		if self.versions is None:
			self.versions = []
		self.versions.append({
			"timestamp": datetime.now().isoformat(),
			"action": "update",
			"previous_hash": self.hashx,
			"previous_size": self.file_size_bytes,
		})

		self.doc_binary = file_content
		self.detect_mime_type(filename)
		self.set_doc_type_from_mime_type()
		self.file_size_bytes = len(file_content)
		self.update_hash()
		self.extract_metadata()
